from pathlib import Path
from datetime import datetime

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

from utils.config import OUTPUT_DIR


def export_to_txt(content: str) -> str:

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = OUTPUT_DIR / f"meeting_minutes_{timestamp}.txt"

    with open(filename, "w", encoding="utf-8") as file:
        file.write(content)

    return str(filename)


def export_to_pdf(content: str) -> str:

    OUTPUT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = OUTPUT_DIR / f"meeting_minutes_{timestamp}.pdf"

    doc = SimpleDocTemplate(str(filename))

    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]

    story = []

    story.append(
        Paragraph(
            content.replace("\n", "<br/>"),
            normal_style,
        )
    )

    doc.build(story)

    return str(filename)


def export_to_docx(content: str) -> str:

    OUTPUT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = OUTPUT_DIR / f"meeting_minutes_{timestamp}.docx"

    document = Document()

    document.add_heading("Meeting Minutes", level=1)

    document.add_paragraph(content)

    document.save(filename)

    return str(filename)


